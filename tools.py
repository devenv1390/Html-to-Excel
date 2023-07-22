from copy import deepcopy

import docx
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from colorama import Fore, Style

special_list = ['4.1', '5.2', '5.3', '5.4', '5.5',
                '5.6', '5.7', '5.8', '5.9', '5.10',
                '5.11', '5.12', '5.13', '5.14',
                '5.15', '5.16', '5.17']


def print_info(text):
    print(Fore.GREEN + text + Style.RESET_ALL)


def print_warning(text):
    print(Fore.YELLOW + text + Style.RESET_ALL)


def print_error(text):
    print(Fore.RED + text + Style.RESET_ALL)


# 填充标题表格
def fill_title_table(table, data_list, doc, file_type):
    print("==========================================")
    print("------ 正在处理标题数据 ------")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if col_index == 2 and row_index < 77:
                flag = 1

                main_text = table.cell(row_index, col_index - 1).text
                num_text = table.cell(row_index, col_index - 2).text
                title_text = num_text + " " + main_text

                for data in data_list:
                    if file_type == 0 or file_type == 1 or file_type == 0.5:
                        if table.cell(row_index, col_index - 2).text == data[1] \
                                and cell.text != 'AUTOSAR网络管理测试' and cell.text != '物理层测试' \
                                and cell.text != '数据链路层测试' and cell.text != '网络管理测试' and cell.text != '应用层测试':

                            if data[1] == '6.1' or data[1] == '6.2' or data[1] == '6.3' or data[1] == '6.4':

                                if (file_type == 1 or file_type == 0.5) and (data[1] == '6.1' or data[1] == '6.2'):
                                    if ('Busoff下NM状态转换' in data[2] or '高负载下的NM状态转换测试' in data[2]) \
                                            and (
                                            main_text == 'BUS-OFF下NM状态转换测试' or main_text == '高负载下的NM状态转换测试'):
                                        compare_set_title_result(cell, data, title_text)
                                        break

                                elif file_type == 0 and data[1] == '6.4':
                                    if '位上升' in data[2] and '下降时间' in data[2] and main_text == '位上升/下降时间':
                                        compare_set_title_result(cell, data, title_text)
                                        break

                                elif data[2] in table.cell(row_index, col_index - 1).text:
                                    compare_set_title_result(cell, data, title_text)
                                    break
                            else:
                                compare_set_title_result(cell, data, title_text)
                                break
                        elif flag == 1 and (cell.text == '' or cell.text == "N/A"):
                            find_text_with_read_table(doc, main_text)
                            # print("Not find but reset: " + title_text)
                            flag = 0
                    else:
                        if table.cell(row_index, col_index - 2).text == data[0] \
                                and cell.text != 'AUTOSAR网络管理测试' and cell.text != '物理层测试' \
                                and cell.text != '数据链路层测试' and cell.text != '网络管理测试' and cell.text != '应用层测试':

                            if data[0] == '6.1' or data[0] == '6.2' or data[0] == '6.3' or data[0] == '6.4':
                                if data[1] in table.cell(row_index, col_index - 1).text:
                                    print("Find and set: " + title_text)
                                    cell.text = data[2]
                                    break
                            else:
                                print("Find and set: " + title_text)
                                cell.text = data[2]

                        elif flag == 1 and (cell.text == '' or cell.text == "N/A"):
                            find_text_with_read_table(doc, title_text)
                            flag = 0

                if cell.text != 'AUTOSAR网络管理测试' and cell.text != '物理层测试' and cell.text != '数据链路层测试' \
                        and cell.text != '网络管理测试' and cell.text != '应用层测试':
                    if cell.text == '':
                        cell.text = "N/A"
                    set_result_type(cell)

    print("------ 完成标题数据处理 ------")
    print("==========================================")


# 判断后填入标题结果
def compare_set_title_result(cell, data, title_text):
    print("Find and set: " + title_text)
    if len(data) > 4:
        if data[4] == 'warning':
            cell.text = "N/A"
        else:
            if data[3] == 'pass':
                cell.text = "OK"
            else:
                cell.text = "NOK"
    else:
        if data[3] == 'pass':
            cell.text = "OK"
        else:
            cell.text = "NOK"


# 查找目标表格的段落并返回索引
def find_paragraph_with_table(doc, table_identifier):
    for i, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            if table_identifier in run.text:
                return i
    return -1


# 复制单元格样式
def copy_cell_style(source_cell, target_cell):
    target_cell.paragraphs[0].clear()
    for paragraph in source_cell.paragraphs:
        new_paragraph = target_cell.add_paragraph()
        new_paragraph.text = paragraph.text
        for run in paragraph.runs:
            new_run = new_paragraph.add_run()
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            new_run.font.color.rgb = run.font.color.rgb


# 将带有相同单元格信息的表格存入一个list中
def find_tables_with_content(doc, target_content):
    tables_with_content = []  # 存储符合条件的表格的列表

    # 遍历文档中的所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 检查目标内容是否在当前单元格的文本中
                if target_content in cell.text:
                    tables_with_content.append(table)
                    break
            else:
                continue
            break

    return tables_with_content


# 拷贝表格并粘贴在下一个位置
def copy_table(document, table_to_copy, target_table_identifier):
    # 获取要复制的表格的行数和列数
    rows = len(table_to_copy.rows)
    cols = len(table_to_copy.columns)

    # 获取要复制的表格的样式
    table_style = table_to_copy.style

    # 创建新表格
    new_table = document.add_table(rows=rows, cols=cols)

    # 复制表格样式
    new_table._element = deepcopy(table_to_copy._element)

    # 设置新表格的样式
    new_table.style = table_style

    # 复制要复制的表格内容到新表格
    for i in range(rows):
        for j in range(cols):
            copy_cell_style(table_to_copy.cell(i, j), new_table.cell(i, j))

    # 查找目标表格的段落
    target_paragraph_index = find_paragraph_with_table(document, target_table_identifier)

    if target_paragraph_index >= 0:
        # 在目标段落的下一个位置插入新表格
        target_paragraph = document.paragraphs[target_paragraph_index + 1]
        new_table_paragraph = target_paragraph.insert_paragraph_before(" ")  # 添加空行
        # new_table_paragraph.add_run().add_break(WD_BREAK.LINE)  # 这也是一种添加空行的方法
        new_table_paragraph._element.getparent().insert(
            new_table_paragraph._element.getparent().index(new_table_paragraph._element) + 1, new_table._element
        )

        tables = document.tables
        if len(tables) > 0:
            # 获取最后一个表格
            last_table = tables[-1]

            # 删除最后一个表格
            parent = last_table._element.getparent()
            parent.remove(last_table._element)
        else:
            print("文档中没有表格。")
    else:
        print("未找到目标表格。")


# 填充普通数据表格
def fill_normal_table(table, data_list):
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if row_index >= 2 and row_index - 2 < len(data_list):
                if col_index == 2:
                    cell.text = str(data_list[row_index - 2][0])

                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中对齐

                    # 遍历单元格内的段落并设置水平居中对齐
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置段落水平居中对齐
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中对齐
                        paragraph.paragraph_format.widow_control = True  # 设置自动换行
                        run = paragraph.runs[0]
                        run.font.size = Pt(10.5)

                        # 遍历段落内的run，并设置字体
                        for run in paragraph.runs:
                            for char in run.text:
                                if 0x4E00 <= ord(char) <= 0x9FFF:
                                    # 中文设置为宋体
                                    run.font.name = '宋体'
                                else:
                                    # 英文、数字和符号设置为Times New Roman
                                    run.font.name = 'Times New Roman'

                elif col_index == 4:
                    cell.text = data_list[row_index - 2][1]
                    if cell.text == 'warning':
                        cell.text = "N/A"
                    else:
                        if cell.text == 'pass':
                            cell.text = "OK"
                        else:
                            cell.text = "NOK"
                    set_result_type(cell)


# 填充特殊数据表格
def fill_special_table(table, data_list):
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if row_index >= 2 and col_index == 4:
                cell.text = "OK"
                for data in data_list:
                    if table.cell(row_index, 3).text in data[0]:
                        cell.text = "NOK"
                        print("Set fail at cell(" + row_index.__str__() + ", 3)")
                        break
                set_result_type(cell)


# 预填充表格
def pre_fill_normal_table(table):
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if row_index >= 2:
                if col_index == 4:
                    cell.text = "N/A"
                    set_result_type(cell)


# 设置测试结果单元格样式
def set_result_type(cell):
    # 设置垂直居中对齐
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    content = cell.text.strip()

    shading_color = None  # 默认为无色
    if content == "OK":
        shading_color = RGBColor(0, 128, 0)  # 绿色
        # print("OK")
    elif content == "NOK":
        shading_color = RGBColor(255, 0, 0)  # 红色
        # print("NOK")
    elif content == "N/A":
        shading_color = RGBColor(255, 255, 255)  # 白色
        # print("N/A")

    # 添加或修改单元格的背景色
    if shading_color is not None:
        if cell._element.tcPr is None:
            cell._element.tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
        shading_element = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
        cell._element.tcPr.append(shading_element)

    # 遍历单元格内的段落并设置水平居中对齐
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.runs[0]
        run.font.size = Pt(10.5)
        run.font.name = 'Times New Roman'


# 读出单元格内容，然后进行预填充
def find_text_with_read_table(doc, target_text):
    paragraphs = doc.paragraphs
    all_tables = doc.tables
    target_text = target_text.encode('utf-8').decode('utf-8')

    for aPara in paragraphs:  # 遍历段落找表格
        if target_text in aPara.text:
            ele = aPara._p.getnext()
            while ele.tag != '' and ele.tag[-3:] != 'tbl':
                ele = ele.getnext()
            if ele.tag != '':
                for table in all_tables:
                    if table._tbl == ele and table.cell(0, 0).text != '测试用例章节':
                        pre_fill_normal_table(table)


# 按标题查找并填充word数据表格
def find_text_with_fill_table(docx_file, target_text, data_list,
                              file_path, index, special_target_result):
    doc = Document(docx_file)

    paragraphs = doc.paragraphs
    all_tables = doc.tables
    target_text = target_text.encode('utf-8').decode('utf-8')

    print("==========================================")
    print("------ 正在处理第" + index.__str__() + "个测试数据 ------")

    for aPara in paragraphs:  # 遍历段落找表格
        if target_text in aPara.text:
            ele = aPara._p.getnext()
            while ele.tag != '' and ele.tag[-3:] != 'tbl':
                ele = ele.getnext()
            if ele.tag != '':
                for table in all_tables:
                    if table._tbl == ele:
                        table.autofit = False
                        print("Find " + target_text)
                        if 'Bus Off恢复时间' in target_text:

                            new_data_list = split_list(data_list, 4)  # 切分数据

                            if len(data_list) > 4:
                                for i in range(1, int(len(data_list) / 4)):
                                    copy_table(doc, table, "8.8 Bus Off恢复时间")

                            # 目标内容，这里是"8.8 Bus Off恢复时间"，可以根据实际需求替换为其他内容
                            target_content = "8.8 Bus Off恢复时间"

                            # 查找带有目标内容的表格
                            tables_with_content = find_tables_with_content(doc, target_content)

                            for idx, _table in enumerate(tables_with_content):
                                fill_normal_table(_table, new_data_list[idx])
                        elif (target_text.split(" ")[0] in special_list) and ("6.4 位上升/下降时间" != target_text):

                            print("Get special title " + target_text)

                            if 'pass' in special_target_result:
                                for row_index, row in enumerate(table.rows):
                                    for col_index, cell in enumerate(row.cells):
                                        if row_index >= 2 and col_index == 4:
                                            cell.text = "OK"
                                            set_result_type(cell)
                            if 'fail' in special_target_result:
                                if target_text.split(" ")[0] == '5.17':
                                    for row_index, row in enumerate(table.rows):
                                        for col_index, cell in enumerate(row.cells):
                                            if row_index >= 2 and col_index == 4:
                                                cell.text = "NOK"
                                                set_result_type(cell)
                                else:
                                    fill_special_table(table, data_list)

                        else:
                            fill_normal_table(table, data_list)

    doc.save(file_path)

    print("------ 完成第" + index.__str__() + "个测试数据的处理 ------")
    print("==========================================")


# 查找并填充word标题表格
def find_text_with_fill_title(docx_file, title_target_text, title_data_list, file_path, file_type):
    doc = Document(docx_file)

    paragraphs = doc.paragraphs
    all_tables = doc.tables

    for aPara in paragraphs:  # 遍历段落找表格
        # print(aPara.text)
        if title_target_text in aPara.text:
            ele = aPara._p.getnext()
            while ele.tag != '' and ele.tag[-3:] != 'tbl':
                ele = ele.getnext()
            if ele.tag != '':
                for table in all_tables:
                    if table._tbl == ele:
                        table.autofit = False
                        fill_title_table(table, title_data_list, doc, file_type)

    doc.save(file_path)


# 检测特定字符串
def find_same_context(cell, target_list):
    # 使用循环遍历找到包含特定字符串的元素
    for _str in target_list:
        if cell.text.strip() in _str[2]:
            # print("find " + cell.text.strip())
            return True
    return False


# 解析HTML中的表格信息
def process_nested_table(table):
    nested_data = []

    rows = table.find_all('tr')

    # 获取表头
    header_row = rows[0]
    headers = []
    for th in header_row.find_all('th'):
        headers.append(th.text.strip())
    # 添加表头到数据
    nested_data.append(headers)

    # 处理表格行
    for row in rows[1:]:
        row_data = []
        for cell in row.find_all(['th', 'td']):
            # 如果单元格内包含其他嵌套表格，则递归处理
            if cell.find('table'):
                nested_table = cell.find('table')
                nested_table_data = process_nested_table(nested_table)
                row_data.append(nested_table_data)
            else:
                row_data.append(cell.get_text().strip())

        # 将行数据添加到嵌套数据中
        nested_data.append(row_data)

    return nested_data


# 解析HTML中的标题信息
def process_table(table):
    data = []

    rows = table.find_all('tr')

    # 处理表格行
    for row in rows[0:]:
        row_data = []
        for cell in row.find_all(['th', 'td']):
            # 如果单元格内包含其他嵌套表格，则递归处理
            if cell.find('table'):
                nested_table = cell.find('table')
                nested_table_data = process_nested_table(nested_table)
                row_data.append(nested_table_data)
            else:
                row_data.append(cell.get_text().strip())

        # 将行数据添加到嵌套数据中
        data.append(row_data)

    return data


# 替换文本
def replace_at_symbol(data_list):
    for data in data_list:
        temp_list = data[2].split("到")
        result = ""
        for item in temp_list:
            result = result + item + "-"
        res = result[:-1]
        data[2] = res


# 去除换行符
def delete_enter(table):
    lines = table.text.split('\n')  # 分割字符串为行

    result = []  # 存储数据的列表

    for line in lines:
        line = line.strip()  # 去除行首尾的空格
        if line:  # 确保行不为空
            result.append(line)  # 将行添加到结果列表

    return result


# 特殊处理2类 HTML 文件的标题表格
def special_duel_with_title(table):
    result = delete_enter(table)
    result_temp = []

    is_done = False  # 判断是否到特殊位置

    for line in result:
        clean_line = line
        if clean_line:  # 确保行不为空
            temp_line = clean_line.split(" ")

            if temp_line[0] == '9.8':  # 最后一项特殊处理，后面的数据可以丢弃
                is_done = True
                clean_line = temp_line[1] + " " + temp_line[2]
                result_temp.append([temp_line[0], clean_line, ''])

            if not is_done and temp_line[0] != '6' and temp_line[0] != '7' \
                    and temp_line[0] != '8' and temp_line[0] != '9':  # 没到特殊位置时进行一般处理

                if 2 <= len(temp_line) <= 4 and temp_line[0] != '7.2' \
                        and temp_line[0] != '9.7' and temp_line[0] != '6.4':  # 一般情况
                    result_temp.append([temp_line[0], temp_line[1], ''])

                elif temp_line[0] == '7.2':  # 该情况需要特殊处理
                    clean_line = temp_line[1] + " " + temp_line[2] + " " + temp_line[3]
                    result_temp.append([temp_line[0], clean_line, ''])

                elif temp_line[0] == '9.7':  # 该情况需要特殊处理
                    clean_line = temp_line[1] + " " + temp_line[2]
                    result_temp.append([temp_line[0], clean_line, ''])

                elif temp_line[0] == '6.4':  # 该情况需要特殊处理
                    if temp_line[1] == '位上升下降时间':
                        temp_line[1] = '位上升/下降时间'
                    else:
                        temp_line[1] = '多网段同步唤醒测试'
                    clean_line = temp_line[1]
                    result_temp.append([temp_line[0], clean_line, ''])

    return result_temp


# 特殊处理2类 HTML 文件的测试数据表格
def special_duel_with_table(table, title_data):
    result = delete_enter(table)
    print(result)
    print("----------")
    i = 0
    for data in result:
        title_data[i][2] = result[1]
        i += 1


# 处理类型0的 HTML 数据，整理后合并成一个final_list
def connect_data_type_zero(table, title_data):
    final_data = []
    flag_count = 0  # 计数，第几个测试项目
    for cell in table:
        for _cell in cell:
            for __cell in _cell:
                if __cell == 'Timestamp':
                    temp_data = cell
                    final_data.append(['序号', '题号', '测试大类项目名称', '大类项目测试结果'])
                    if len(title_data[flag_count]) > 4:
                        title_data[flag_count].pop()
                        final_data.append(title_data[flag_count])
                    else:
                        final_data.append(title_data[flag_count])
                    flag_count += 1
                    final_data.append(['测试项目', '测试标准', '测试数值', '测试结果'])
                    for i in range(4, len(temp_data) + 1, 4):
                        # print(temp_data)
                        temp_cell_data = temp_data[i]
                        temp_cell_data.append(temp_data[i - 3].pop())
                        final_data.append(temp_cell_data)
                    final_data.append([' '])
                    final_data.append([' '])
    return final_data


# 处理类型1和类型0.5的 HTML 数据，整理后合并成一个final_list
def connect_data_type_one(table, title_data):
    final_data = []
    flag_count = 0  # 计数，第几个测试项目
    skip_index = 2
    for cell_index, cell in enumerate(table):
        for _cell_index, _cell in enumerate(cell):
            for __cell in _cell:
                if __cell == 'Timestamp':
                    is_special = 0
                    # print(table[cell_index - 3][0][1])
                    temp_data = cell
                    temp_title = table[cell_index - 3][0][1]
                    for result in special_list:

                        if result in temp_title and 'CAN总线电压' not in temp_title \
                                and 'CAN_H与CAN_L的内阻' not in temp_title \
                                and 'CAN_H与CAN_L之间的差分电阻' not in temp_title \
                                and '位上升/下降时间' not in temp_title:
                            # print("special!")
                            final_data.append(['序号', '题号', '测试大类项目名称', '大类项目测试结果'])
                            if len(title_data[flag_count]) > 4:
                                title_data[flag_count].pop()
                                final_data.append(title_data[flag_count])
                            else:
                                final_data.append(title_data[flag_count])
                            flag_count += 1
                            if title_data[flag_count - 1][1] == '':
                                final_data.pop()
                                temp_title_data, flag_count = next_title_context(title_data, flag_count)
                                final_data.append(temp_title_data)
                                flag_count += 1
                            # print(final_data[len(final_data) - 1])
                            final_data.append(['测试项目', '测试标准', '测试数值', '测试结果'])
                            local = []
                            network = []
                            diagnosis = []
                            data = []
                            tip_list = ['本地唤醒', 'NM唤醒', '诊断唤醒', '数据帧唤醒']
                            pre_data = ""
                            # print(temp_data)
                            for i in range(2, len(temp_data) + 1, 4):
                                # print(temp_data[i])
                                if 'DUT无应用报文发送' in temp_data[i][1]:
                                    for tip_index, tip in enumerate(tip_list):
                                        if tip in pre_data:
                                            if tip_index == 0:
                                                local.append(temp_data[i - 1].pop())
                                            elif tip_index == 1:
                                                network.append(temp_data[i - 1].pop())
                                            elif tip_index == 2:
                                                diagnosis.append(temp_data[i - 1].pop())
                                            else:
                                                data.append(temp_data[i - 1].pop())
                                if '本地唤醒' in temp_data[i][1]:
                                    local.append(temp_data[i - 1].pop())
                                if 'NM唤醒' in temp_data[i][1] or 'NM报文唤醒' in temp_data[i][1]:
                                    network.append(temp_data[i - 1].pop())
                                if '诊断唤醒' in temp_data[i][1]:
                                    diagnosis.append(temp_data[i - 1].pop())
                                if '数据帧唤醒' in temp_data[i][1]:
                                    data.append(temp_data[i - 1].pop())
                                pre_data = temp_data[i][1]

                            for item in local:
                                if item == 'fail':
                                    final_data.append(['', '', '本地事件', 'fail'])
                                    break
                            for item in network:
                                if item == 'fail':
                                    final_data.append(['', '', '网络管理报文', 'fail'])
                                    break
                            for item in diagnosis:
                                if item == 'fail':
                                    final_data.append(['', '', '诊断报文', 'fail'])
                                    break
                            for item in data:
                                if item == 'fail':
                                    final_data.append(['', '', '数据帧报文', 'fail'])
                                    break
                            # print("---------------")
                            final_data.append([' '])
                            final_data.append([' '])
                            is_special = 1
                            break
                    if is_special == 1:
                        break
                    if skip_index > 0:
                        skip_index -= 1
                    else:
                        final_data.append(['序号', '题号', '测试大类项目名称', '大类项目测试结果'])
                        if len(title_data[flag_count]) > 4:
                            title_data[flag_count].pop()
                            final_data.append(title_data[flag_count])
                        else:
                            final_data.append(title_data[flag_count])
                        flag_count += 1
                        if title_data[flag_count - 1][1] == '':
                            final_data.pop()
                            temp_title_data, flag_count = next_title_context(title_data, flag_count)
                            final_data.append(temp_title_data)
                            flag_count += 1
                        # print(final_data[len(final_data) - 1])
                        final_data.append(['测试项目', '测试标准', '测试数值', '测试结果'])
                        for i in range(4, len(temp_data) + 1, 4):
                            temp_cell_data = temp_data[i]
                            temp_cell_data.append(temp_data[i - 3].pop())
                            final_data.append(temp_cell_data)
                            # print(temp_cell_data)
                        # print("---------------")
                        final_data.append([' '])
                        final_data.append([' '])
    return final_data


# 查询标题表格的下一个内容
def next_title_context(title_data, flag_count):
    if title_data[flag_count][1] == '':
        return next_title_context(title_data, flag_count + 1)
    else:
        if len(title_data[flag_count]) > 4:
            title_data[flag_count].pop()
        return title_data[flag_count], flag_count


# 提取final_list中的单个list元素作为data_list填入表格
def get_list_from_final(final_list):
    data_list = []

    for data_index, data in enumerate(final_list):
        temp_list = []
        if len(data) > 2:
            if data[1] == '题号':
                if final_list[data_index + 1][1] in special_list:
                    temp_list.append([final_list[data_index + 1][1] + " " + final_list[data_index + 1][2] +
                                      final_list[data_index + 1][3]])
                else:
                    temp_list.append([final_list[data_index + 1][1] + " " + final_list[data_index + 1][2]])
                for _data in range(data_index + 3, len(final_list) - 1):
                    if final_list[_data] == [' ']:
                        break
                    else:
                        temp_list.append([final_list[_data][2], final_list[_data][3]])
                data_list.append(temp_list)
    return data_list


# 把一个等长的列表切分成若干个等长的子列表
def split_list(lst, chunk_size):
    return [lst[i:i + chunk_size] for i in range(0, len(lst), chunk_size)]


# list内元素计数
def count_element(lst, target):
    count = 0
    for item in lst:
        if isinstance(item, list):
            count += count_element(item, target)  # 递归调用处理嵌套列表
        elif item == target:
            count += 1
    return count

# #将表格填充到EXCEL中的cell
# def fill_cell(ws):
#     prev_row_values = []  # 用于存储上一行的值
#     for row in ws.rows:
#         for cell in row:
#             cell.alignment = alignment
#             if cell.value == 'pass':
#                 cell.fill = green_fill
#                 cell.font = data_font
#                 cell.value = 'OK'
#                 cell.border = thin_border
#             elif cell.value == 'warning':
#                 cell.fill = yellow_fill
#                 cell.font = data_font
#                 cell.value = 'N/A'
#                 cell.border = thin_border
#             elif cell.value == 'fail':
#                 cell.fill = red_fill
#                 cell.font = data_font
#                 cell.value = 'NOK'
#                 cell.border = thin_border
#             elif cell.value is not None:
#                 if cell.value != ' ':
#                     cell.border = thin_border
#                     cell.font = value_font
#                     for _i in header_list:
#                         if cell.value == _i:
#                             cell.font = header_font
#                             cell.fill = blue_fill
#                             cell.border = thin_border
#                     for header_value in ['序号', '题号', '测试大类项目名称', '大类项目测试结果']:
#                         if header_value in prev_row_values:
#                             cell.font = header_font
#         # 更新上一行的值
#         prev_row_values = [cell.value for cell in row]
#
# # 创建一个新的 DataFrame 来存储嵌套数据
#     df = pd.DataFrame(final_list)
#
#     # 写入 Excel 文件
#     df.to_excel(output_filepath, index=False)
#
#     # 修改 Excel 文件格式
#     wb = openpyxl.load_workbook(output_filepath)
#
#     # 获取默认的活动工作表
#     ws_detail = wb.active
#     ws_detail.title = "详细信息"
#
#     # 创建“总览”工作簿
#     ws_overview = wb.create_sheet(title="总览")
#     for i in range(len(title_data)):
#         for j in range(len(title_data[i])):
#             if isinstance(title_data[i][j], list):
#                 ws_overview.append(title_data[i][j])
#             else:
#                 ws_overview.cell(i + 1, j + 1).value = title_data[i][j]  # 写入数据
#                 ws_overview.cell(i + 1, j + 1).alignment = Alignment(horizontal='center', vertical='center')  # 居中对齐
#
#     # 设置列宽
#     ws_detail_column_widths = [20, 60, 70, 10, 10]
#     for i, width in enumerate(ws_detail_column_widths):
#         col_letter = chr(65 + i)  # A, B, C...
#         ws_detail.column_dimensions[col_letter].width = width
#
#     ws_overview_column_widths = [26, 10, 26, 10]
#     for i, width in enumerate(ws_overview_column_widths):
#         col_letter = chr(65 + i)  # A, B, C...
#         ws_overview.column_dimensions[col_letter].width = width
#
#     # 定义单元格边框样式
#     thin_border = Border(left=Side(style='thin'),
#                          right=Side(style='thin'),
#                          top=Side(style='thin'),
#                          bottom=Side(style='thin'))
#
#     # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
#     alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
#
#     # 定义单元格样式
#     header_font = Font(bold=True, size=14)
#     data_font = Font(bold=True, size=14)
#     value_font = Font(size=14)
#     yellow_fill = PatternFill(fill_type='solid', fgColor='FFFF00')
#     green_fill = PatternFill(fill_type='solid', fgColor='00FF00')
#     red_fill = PatternFill(fill_type='solid', fgColor='FF0000')
#     blue_fill = PatternFill(fill_type='solid', fgColor='00B0F0')
#
#     # 拟造条件list
#     header_list = ['测试项目', '测试标准', '测试数值', '测试结果', '序号', '题号', '测试大类项目名称',
#                    '大类项目测试结果']
#     # 填充表格数据
#     fill_cell(ws_overview)
#     fill_cell(ws_detail)
#
#     temp_cell1 = ws_overview["c4"]
#     temp_cell2 = ws_overview["c5"]
#
#     temp_cell1.fill = green_fill
#     temp_cell2.fill = red_fill
#     temp_cell1.font = data_font
#     temp_cell2.font = data_font
#
#     # 将工作表名称排序
#     sort_lst = sorted(wb.sheetnames)
#
#     # 重新排列工作表
#     for sheet_name in sort_lst:
#         ws = wb[sheet_name]
#         wb.remove(ws)
#         wb._add_sheet(ws)
#
#     # 保存 Excel 文件
#     wb.save(output_filepath)
