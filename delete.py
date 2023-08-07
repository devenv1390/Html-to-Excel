import os

from docx import Document

special_list = ['AUTOSAR网络管理测试', '物理层测试', '数据链路层测试',
                '网络管理测试', '应用层测试', 'CAN总线电压',
                '故障管理', '通信电压', 'CAN_H与CAN_L的内阻']


def remove_tables_from_docx(input_file_path, output_file_path):
    try:
        # 加载输入的docx文件
        doc = Document(input_file_path)

        title_data = find_title_table(doc)
        # print(title_data)

        for target in title_data:
            find_and_delete_table(doc, target)
            print(f"已删除表格 {target}")

        # 将修改后的文档保存到输出文件
        doc.save(output_file_path)

        print(" ")
        print("当前文件处理完成")
        print("==================")
    except Exception as e:
        print("发生错误，报错信息为:", e)


# 查找并填充word标题表格
def find_title_table(doc):
    title_data = []

    paragraphs = doc.paragraphs
    all_tables = doc.tables

    for aPara in paragraphs:  # 遍历段落找表格
        # print(aPara.text)
        if "测试项目总览" in aPara.text:
            ele = aPara._p.getnext()
            while ele.tag != '' and ele.tag[-3:] != 'tbl':
                ele = ele.getnext()
            if ele.tag != '':
                for table in all_tables:
                    if table._tbl == ele:
                        for row_index, row in enumerate(table.rows):
                            for col_index, cell in enumerate(row.cells):
                                if col_index == 1 and 1 < row_index < 77:
                                    temp_cell = table.cell(row_index, col_index + 1)
                                    if cell.text not in special_list:
                                        if temp_cell.text == "N/A":
                                            title_data.append(cell.text)
    return title_data


# 查找并删除文档中的表格
def find_and_delete_table(doc, target_text):
    paragraphs = doc.paragraphs
    all_tables = doc.tables
    target_text = target_text.encode('utf-8').decode('utf-8')

    for aPara in paragraphs:
        if target_text in aPara.text:
            ele = aPara._p.getnext()
            while ele is not None and ele.tag != '' and ele.tag[-3:] != 'tbl':
                ele = ele.getnext()
            if ele is not None and ele.tag != '':
                for table in all_tables:
                    if table._tbl == ele and table.cell(0, 0).text != '测试用例章节':
                        table._element.getparent().remove(table._element)
                        return



if __name__ == "__main__":
    docx_file = os.listdir('output')
    # 在所有 docx 文件中遍历处理
    for i, filename in enumerate(docx_file, start=1):
        print(f"正在处理{filename}")
        input_file_path = os.path.join('output', filename)
        output_file_path = os.path.join('output', filename)
        remove_tables_from_docx(input_file_path, output_file_path)
    os.system("pause")
