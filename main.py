# import os
#
# import openpyxl
# import pandas as pd
# from bs4 import BeautifulSoup
# from openpyxl.styles import Border, Side
# from openpyxl.styles import Font, Alignment
# # excel 数据样式设置类
# from openpyxl.styles import PatternFill
#
# # 创建 output 文件夹（如果不存在）
# if not os.path.exists('output'):
#     os.makedirs('output')
#
#
# # 解析HTML中的表格信息
# def process_nested_table(table):
#     nested_data = []
#
#     rows = table.find_all('tr')
#
#     # 获取表头
#     header_row = rows[0]
#     headers = []
#     for th in header_row.find_all('th'):
#         headers.append(th.text.strip())
#     # 添加表头到数据
#     nested_data.append(headers)
#
#     # 处理表格行
#     for row in rows[1:]:
#         row_data = []
#         for cell in row.find_all(['th', 'td']):
#             # 如果单元格内包含其他嵌套表格，则递归处理
#             if cell.find('table'):
#                 nested_table = cell.find('table')
#                 nested_table_data = process_nested_table(nested_table)
#                 row_data.append(nested_table_data)
#             else:
#                 row_data.append(cell.get_text().strip())
#
#         # 将行数据添加到嵌套数据中
#         nested_data.append(row_data)
#
#     return nested_data
#
#
# def process_table(table):
#     data = []
#
#     rows = table.find_all('tr')
#
#     # 处理表格行
#     for row in rows[0:]:
#         row_data = []
#         for cell in row.find_all(['th', 'td']):
#             # 如果单元格内包含其他嵌套表格，则递归处理
#             if cell.find('table'):
#                 nested_table = cell.find('table')
#                 nested_table_data = process_nested_table(nested_table)
#                 row_data.append(nested_table_data)
#             else:
#                 row_data.append(cell.get_text().strip())
#
#         # 将行数据添加到嵌套数据中
#         data.append(row_data)
#
#     return data
#
#
# def replace_at_symbol(lst):
#     result = []
#     for item in lst:
#         if isinstance(item, list):
#             result.append(replace_at_symbol(item))
#         elif isinstance(item, str) and item == '<br>' and item == '</br>':
#             result.append(';')
#         else:
#             result.append(item)
#     return result
#
#
# # 处理数据，合并成一个list
# def connect_data(table):
#     final_data = []
#     title_data = table[9]  # 测试项的标题
#     title_data.insert(1, ['1', 'null', 'DutSelfCheck', 'null'])  # 自检这一项读不出来，要手动初始化
#     flag_count = 1  # 计数，第几个测试项目
#     for cell in table:
#         for _cell in cell:
#             for __cell in _cell:
#                 if __cell == 'Timestamp':
#                     temp_data = cell
#                     final_data.append(['序号', '题号', '测试大类项目名称', '大类项目测试结果'])
#                     if len(title_data[flag_count]) > 4:
#                         title_data[flag_count].pop()
#                         final_data.append(title_data[flag_count])
#                     else:
#                         final_data.append(title_data[flag_count])
#                     flag_count += 1
#                     final_data.append(['测试项目', '测试标准', '测试数值', '测试结果'])
#                     for i in range(4, len(temp_data) + 1, 4):
#                         temp_cell_data = temp_data[i]
#                         temp_cell_data.append(temp_data[i - 3].pop())
#                         final_data.append(temp_cell_data)
#                     final_data.append([' '])
#                     final_data.append([' '])
#     return final_data
#
# #list内元素计数
# def count_element(lst, target):
#     count = 0
#     for item in lst:
#         if isinstance(item, list):
#             count += count_element(item, target)  # 递归调用处理嵌套列表
#         elif item == target:
#             count += 1
#     return count
#
# #填充cell
# def fill_cell(ws):
#     prev_row_values = []  # 用于存储上一行的值
#     for row in ws.rows:
#         for cell in row:
#             cell.alignment = alignment
#             if cell.value == 'pass':
#                 cell.fill = green_fill
#                 cell.font = data_font
#                 cell.value = 'OK'
#                 cell.border = thin_border
#             elif cell.value == 'warning':
#                 cell.fill = yellow_fill
#                 cell.font = data_font
#                 cell.value = 'N/A'
#                 cell.border = thin_border
#             elif cell.value == 'fail':
#                 cell.fill = red_fill
#                 cell.font = data_font
#                 cell.value = 'NOK'
#                 cell.border = thin_border
#             elif cell.value is not None:
#                 if cell.value != ' ':
#                     cell.border = thin_border
#                     cell.font = value_font
#                     for _i in header_list:
#                         if cell.value == _i:
#                             cell.font = header_font
#                             cell.fill = blue_fill
#                             cell.border = thin_border
#                     for header_value in ['序号', '题号', '测试大类项目名称', '大类项目测试结果']:
#                         if header_value in prev_row_values:
#                             cell.font = header_font
#         # 更新上一行的值
#         prev_row_values = [cell.value for cell in row]
#
#
# # 获取 input 文件夹下的所有 HTML 文件
# html_files = os.listdir('input')
# print("------ 共有" + len(html_files).__str__() + "个文件待处理 ------")
# i = 1
# for filename in html_files:
#     print("------ 正在处理第" + i.__str__() + "个文件 ------")
#     i += 1
#     # 拼接文件路径
#     input_filepath = os.path.join('input', filename)
#     output_filepath = os.path.join('output', filename.replace('.html', '.xlsx'))
#
#     # 读取 HTML 文件
#     with open(input_filepath, 'r', encoding='utf-8') as file:
#         html_content = file.read()
#
#     # 使用 BeautifulSoup 解析 HTML
#     soup = BeautifulSoup(html_content, 'html.parser')
#
#     # 处理多级嵌套表格数据
#     tables = soup.find_all('table')
#     nested_tables_data = []
#     nested_temp_data = []
#     for table in tables:
#         table_data = process_nested_table(table)
#         nested_tables_data.append(table_data)
#         temp_data = process_table(table)
#         nested_temp_data.append(temp_data)
#
#     # 算通过、不通过、没测试
#     total = len(nested_temp_data[9])
#     warning_number = count_element(nested_temp_data[9], 'warning')
#     fail_number = count_element(nested_temp_data[9], 'fail')
#     pass_number = total - fail_number - warning_number
#     complete_number = pass_number + fail_number
#
#     per_complete = round(complete_number / total, 2) * 100
#     per_not_complete = 100 - per_complete
#     per_pass = round(pass_number / total, 2) * 100
#     per_fail = round(fail_number / total, 2) * 100
#
#     str_pre_complete = per_complete.__str__() + "%"
#     str_per_not_complete = per_not_complete.__str__() + "%"
#     str_per_pass = per_pass.__str__() + "%"
#     str_per_fail = per_fail.__str__() + "%"
#
#     title_data = [['测试项目总数', total, ''], ['已完成的测试项目', complete_number, "完成率" + str_pre_complete],
#                   ['未完成的测试', warning_number, "未完成率" + str_per_not_complete],
#                   ['通过的测试数量', pass_number, "通过率" + str_per_pass],
#                   ['未通过的测试数量', fail_number, "未通过率" + str_per_fail], [' '], nested_temp_data[9]]
#
#     # 新建一个list存储修改的数据
#     final_list = connect_data(nested_tables_data)
#     final_list[1] = title_data[6][0]
#     final_list = replace_at_symbol(final_list)
#
#     # 创建一个新的 DataFrame 来存储嵌套数据
#     df = pd.DataFrame(final_list)
#
#     # 写入 Excel 文件
#     df.to_excel(output_filepath, index=False)
#
#     # 修改 Excel 文件格式
#     wb = openpyxl.load_workbook(output_filepath)
#
#     # 获取默认的活动工作表
#     ws_detail = wb.active
#     ws_detail.title = "详细信息"
#
#     # 创建“总览”工作簿
#     ws_overview = wb.create_sheet(title="总览")
#     for i in range(len(title_data)):
#         for j in range(len(title_data[i])):
#             if isinstance(title_data[i][j], list):
#                 ws_overview.append(title_data[i][j])
#             else:
#                 ws_overview.cell(i + 1, j + 1).value = title_data[i][j]  # 写入数据
#                 ws_overview.cell(i + 1, j + 1).alignment = Alignment(horizontal='center', vertical='center')  # 居中对齐
#
#     # 设置列宽
#     ws_detail_column_widths = [20, 60, 70, 10, 10]
#     for i, width in enumerate(ws_detail_column_widths):
#         col_letter = chr(65 + i)  # A, B, C...
#         ws_detail.column_dimensions[col_letter].width = width
#
#     ws_overview_column_widths = [26, 10, 26, 10]
#     for i, width in enumerate(ws_overview_column_widths):
#         col_letter = chr(65 + i)  # A, B, C...
#         ws_overview.column_dimensions[col_letter].width = width
#
#     # 定义单元格边框样式
#     thin_border = Border(left=Side(style='thin'),
#                          right=Side(style='thin'),
#                          top=Side(style='thin'),
#                          bottom=Side(style='thin'))
#
#     # 设置单元格对齐方式 Alignment(horizontal=水平对齐模式,vertical=垂直对齐模式,text_rotation=旋转角度,wrap_text=是否自动换行)
#     alignment = Alignment(horizontal='center', vertical='center', text_rotation=0, wrap_text=True)
#
#     # 定义单元格样式
#     header_font = Font(bold=True, size=14)
#     data_font = Font(bold=True, size=14)
#     value_font = Font(size=14)
#     yellow_fill = PatternFill(fill_type='solid', fgColor='FFFF00')
#     green_fill = PatternFill(fill_type='solid', fgColor='00FF00')
#     red_fill = PatternFill(fill_type='solid', fgColor='FF0000')
#     blue_fill = PatternFill(fill_type='solid', fgColor='00B0F0')
#
#     # 拟造条件list
#     header_list = ['测试项目', '测试标准', '测试数值', '测试结果', '序号', '题号', '测试大类项目名称',
#                    '大类项目测试结果']
#     # 填充表格数据
#     fill_cell(ws_overview)
#     fill_cell(ws_detail)
#
#     temp_cell1 = ws_overview["c4"]
#     temp_cell2 = ws_overview["c5"]
#
#     temp_cell1.fill = green_fill
#     temp_cell2.fill = red_fill
#     temp_cell1.font = data_font
#     temp_cell2.font = data_font
#
#     # 将工作表名称排序
#     sort_lst = sorted(wb.sheetnames)
#
#     # 重新排列工作表
#     for sheet_name in sort_lst:
#         ws = wb[sheet_name]
#         wb.remove(ws)
#         wb._add_sheet(ws)
#
#     # 保存 Excel 文件
#     wb.save(output_filepath)
#
# print("------ 已处理完成所有文件 ------")
# # os.system("pause")
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


def find_text_with_fill_table(docx_file, target_text, data_list):
    doc = Document(docx_file)
    found_text = False

    for paragraph in doc.paragraphs:
        if target_text in paragraph.text:
            found_text = True
        elif found_text:
            print(f"找到匹配的文本 '{target_text}' 在段落: {paragraph.text}")
            found_text = False

    for table in doc.tables:
        table.autofit = False
        if found_text:
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    if row_index >= 2 and row_index - 2 < len(data_list):
                        if col_index == 2:

                            cell.width = Cm(5.28)  # 设置单元格宽度为 5 厘米
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中对齐

                            # 遍历单元格内的段落并设置水平居中对齐
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置段落水平居中对齐
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 设置单元格垂直居中对齐
                                paragraph.paragraph_format.widow_control = True  # 设置自动换行
                                run = paragraph.runs[0]
                                run.font.size = Pt(10.5)

                            cell.text = str(data_list[row_index - 2][0])

                            # 遍历段落内的run，并设置字体
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    for char in run.text:
                                        if 0x4E00 <= ord(char) <= 0x9FFF:
                                            # 中文设置为宋体
                                            run.font.name = '宋体'
                                        else:
                                            # 英文、数字和符号设置为Times New Roman
                                            run.font.name = 'Times New Roman'

                        elif col_index == 4:
                            cell.text = data_list[row_index - 2][1]

                            # 设置垂直居中对齐
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            content = cell.text.strip()

                            shading_color = None  # 默认为无色
                            if content == "OK":
                                shading_color = RGBColor(0, 255, 0)  # 绿色
                            elif content == "NOK":
                                shading_color = RGBColor(255, 0, 0)  # 红色

                            # 添加或修改单元格的背景色
                            if shading_color is not None:
                                if cell._element.tcPr is None:
                                    cell._element.tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                                shading_element = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
                                cell._element.tcPr.append(shading_element)

                            # 遍历单元格内的段落并设置水平居中对齐
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                run = paragraph.runs[0]
                                run.font.size = Pt(10.5)
                                run.font.name = 'Times New Roman'

                    print(cell.text)
            break

        for row in table.rows:
            for cell in row.cells:
                if target_text in cell.text:
                    found_text = True
                    break
            if found_text:
                break

    doc.save("filled_document.docx")


# 指定要搜索的文本、文档路径和数据列表
target_text = "6.1.2"
docx_file = "template.docx"
column_widths = [1.76, 3.53, 5.29, 3.53, 2.47]
data_list = [["哈哈哈哈3.51V", "OK"], ["1.51V", "NOK"], ["2.00V", "N/A"], ["2.51V", "NOK"]]

# 执行搜索并填充表格
find_text_with_fill_table(docx_file, target_text, data_list)
